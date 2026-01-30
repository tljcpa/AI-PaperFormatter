import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Tuple

from app.models.schema import DocumentDSL, ContentType, FontStyle, Alignment

class DocxRenderer:
    """
    物理渲染引擎：将 DocumentDSL 数据结构转换为二进制 .docx 文件。
    """

    def __init__(self):
        pass

    def _hex_to_rgb(self, hex_color: str) -> Tuple[int, int, int]:
        """将 Hex (#000000) 转为 RGB Tuple"""
        if not hex_color:
            return (0, 0, 0)
        hex_color = hex_color.lstrip('#')
        try:
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        except:
            return (0, 0, 0)

    def _apply_paragraph_format(self, paragraph, style: FontStyle):
        """应用段落级样式 (对齐、缩进、行距)"""
        if not style:
            return

        # 1. 对齐
        if style.align == Alignment.CENTER:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style.align == Alignment.RIGHT:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif style.align == Alignment.JUSTIFY:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif style.align == Alignment.LEFT:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 2. 行距
        if style.line_spacing:
            paragraph.paragraph_format.line_spacing = style.line_spacing
        
        # 3. 段间距
        if style.space_before:
            paragraph.paragraph_format.space_before = Pt(style.space_before)
        if style.space_after:
            paragraph.paragraph_format.space_after = Pt(style.space_after)

    def _apply_run_format(self, run, style: FontStyle):
        """应用字符级样式 (字体、大小、颜色、粗体)"""
        if not style:
            return

        # 1. 字体大小
        if style.size:
            run.font.size = Pt(style.size)

        # 2. 粗体/斜体
        if style.bold is not None:
            run.font.bold = style.bold
        if style.italic is not None:
            run.font.italic = style.italic

        # 3. 颜色
        if style.color:
            run.font.color.rgb = RGBColor(*self._hex_to_rgb(style.color))

        # 4. 字体设置 (核心：中西文兼容)
        if style.family:
            run.font.name = style.family
            # 强制设置东亚字体 (针对中文)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), style.family)

    def render(self, dsl: DocumentDSL, output_path: str) -> str:
        """
        渲染入口函数。
        """
        doc = Document()

        # 遍历 DSL 中的每一个内容块
        for block in dsl.content_blocks:
            
            # --- 确定当前块使用的样式 ---
            # 逻辑：Global Style < Override Style
            
            # 1. 获取该类型的全局默认样式
            base_style = getattr(dsl.style_config, block.type.value, None)
            
            # 2. 检查是否有块级覆盖 (User Override specific block)
            active_style = base_style
            if block.style_override:
                # 这里做一个简单的属性覆盖合并 (浅层)
                # 实际生产中可以使用 MergerEngine._deep_update
                if active_style:
                    # 复制一份以免修改全局配置
                    active_style = active_style.model_copy(update=block.style_override.model_dump(exclude_unset=True))
                else:
                    active_style = block.style_override

            # --- 执行渲染 ---
            if block.type in [ContentType.HEADING_1, ContentType.HEADING_2, ContentType.HEADING_3, ContentType.BODY_TEXT, ContentType.CAPTION]:
                # 创建段落
                p = doc.add_paragraph()
                # 应用段落样式
                if active_style:
                    self._apply_paragraph_format(p, active_style)
                
                # 添加文字 Run
                run = p.add_run(block.text)
                # 应用文字样式
                if active_style:
                    self._apply_run_format(run, active_style)
            
            elif block.type == ContentType.IMAGE_HOOK:
                # 简单处理图片占位符
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[图片占位符: {block.source_data}]")
                run.font.color.rgb = RGBColor(255, 0, 0)
            
            # TODO: 实现 Table Hook

        # 保存文件
        doc.save(output_path)
        return output_path

# 单例导出
renderer = DocxRenderer() 
